"""Sinh file Word cho mục 3.3 Thiết kế cơ sở dữ liệu."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, size=13, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 14, 2: 13, 3: 13}
    set_font(run, size=sizes.get(level, 13), bold=True)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.27) if indent else Cm(0)
    run = p.add_run(text)
    set_font(run, size=13)
    return p


def add_table(doc, caption, headers, rows):
    # Caption
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after = Pt(4)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption)
    set_font(cr, size=13, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=12, bold=True)
        cell._tc.get_or_add_tcPr().append(
            OxmlElement("w:shd")
        )

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cell.paragraphs[0].add_run(str(val))
            set_font(r, size=12)

    doc.add_paragraph()


doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ══════════════════════════════════════════════════════
# 3.3 THIẾT KẾ CƠ SỞ DỮ LIỆU
# ══════════════════════════════════════════════════════
add_heading(doc, "3.3. Thiết kế cơ sở dữ liệu", level=1)
add_body(doc, (
    "Hệ thống NomNaSite sử dụng chiến lược lưu trữ lai (hybrid storage) với ba lớp "
    "dữ liệu hoạt động song song: SQLite local lưu trữ từ điển và dữ liệu cục bộ; "
    "Firebase Authentication và Realtime Database quản lý tài khoản, phân quyền trên "
    "cloud; Supabase Storage lưu trữ ảnh tài liệu OCR. Thiết kế này đảm bảo hệ thống "
    "hoạt động ổn định cả khi có lẫn khi mất kết nối Internet."
), indent=True)

# ── 3.3.1 Tổng quan ──────────────────────────────────
add_heading(doc, "3.3.1. Tổng quan kiến trúc lưu trữ", level=2)
add_body(doc, (
    "Hệ thống áp dụng chiến lược kết hợp: ràng buộc khóa ngoại (FK) được thực thi ở "
    "tầng cơ sở dữ liệu cho các quan hệ nội bộ trong SQLite (ocr_sessions → ocr_boxes), "
    "trong khi các tham chiếu xuyên hệ thống — như liên kết giữa username trong SQLite "
    "và tài khoản Firebase Authentication — được kiểm soát ở tầng ứng dụng. Đây là "
    "phương án tối ưu cho kiến trúc hybrid vì không thể tạo FK cứng giữa hai hệ thống "
    "lưu trữ khác nhau."
), indent=True)

add_body(doc, "Ba lớp lưu trữ chính:")
add_body(doc, "•  SQLite (dictionary.db): Lưu từ điển Hán Nôm, phiên OCR, lịch sử dịch, bộ nhớ AI — hoạt động hoàn toàn offline.")
add_body(doc, "•  Firebase Cloud: Authentication quản lý tài khoản; Realtime Database lưu phân quyền (admin/user).")
add_body(doc, "•  Supabase Cloud: Storage bucket lưu ảnh tài liệu OCR; đồng bộ phiên OCR khi có kết nối.")

# ── 3.3.2 SQLite ─────────────────────────────────────
add_heading(doc, "3.3.2. Cơ sở dữ liệu SQLite (dictionary.db)", level=2)
add_body(doc, (
    "Hệ thống sử dụng một file SQLite duy nhất dictionary.db, được khởi tạo tự động "
    "khi ứng dụng chạy lần đầu qua hàm init_database() trong handler/init_database.py. "
    "File này chứa 7 bảng chia thành 3 nhóm chức năng."
), indent=True)

# 3.3.2.1 Từ điển
add_heading(doc, "3.3.2.1. Nhóm bảng Từ điển", level=3)
add_body(doc, (
    "Nhóm này gồm 3 bảng phục vụ chức năng tra cứu và dịch thuật, hoạt động hoàn toàn "
    "độc lập và không có quan hệ FK với nhóm bảng khác."
), indent=True)

add_table(doc,
    "Bảng 3.1: Bảng translations — Từ điển chính Hán Nôm (19.880 entries)",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",          "INTEGER",   "PK, AUTOINCREMENT",  "Khóa chính tự tăng"),
        ("vietnamese",  "TEXT",      "",                   "Từ Quốc Ngữ"),
        ("han_nom",     "TEXT",      "",                   "Chữ Hán Nôm tương ứng"),
        ("han_viet",    "TEXT",      "",                   "Phiên âm Hán Việt"),
        ("meaning",     "TEXT",      "",                   "Nghĩa giải thích"),
    ]
)

add_table(doc,
    "Bảng 3.2: Bảng ai_translations — Cache kết quả dịch AI (29.430 entries)",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",          "INTEGER",   "PK, AUTOINCREMENT",  "Khóa chính tự tăng"),
        ("nom_text",    "TEXT",      "UNIQUE",             "Văn bản Hán Nôm đầu vào"),
        ("meaning",     "TEXT",      "",                   "Phiên âm Hán Việt"),
        ("poetry",      "TEXT",      "",                   "Nghĩa thơ văn"),
        ("category",    "TEXT",      "",                   "Phân loại từ"),
        ("source",      "TEXT",      "",                   "Nguồn dữ liệu"),
        ("vi_meaning",  "TEXT",      "",                   "Nghĩa tiếng Việt"),
        ("created_at",  "TIMESTAMP", "DEFAULT NOW",        "Thời gian tạo"),
    ]
)

add_table(doc,
    "Bảng 3.3: Bảng ai_memory — Bộ nhớ AI cache tái sử dụng",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",                  "INTEGER", "PK, AUTOINCREMENT", "Khóa chính tự tăng"),
        ("nom_text",            "TEXT",    "",                  "Văn bản Hán Nôm"),
        ("modern_meaning",      "TEXT",    "",                  "Nghĩa hiện đại"),
        ("poetic_translation",  "TEXT",    "",                  "Nghĩa thơ văn"),
        ("usage_count",         "INTEGER", "DEFAULT 1",         "Số lần được dùng"),
    ]
)

# 3.3.2.2 Phiên OCR
add_heading(doc, "3.3.2.2. Nhóm bảng Phiên OCR", level=3)
add_body(doc, (
    "Nhóm này gồm 2 bảng có quan hệ khóa ngoại (FK) duy nhất trong hệ thống: "
    "mỗi phiên OCR (ocr_sessions) chứa nhiều vùng chữ (ocr_boxes) theo quan hệ 1–N. "
    "Khi xóa một phiên, tất cả các vùng chữ liên quan cũng bị xóa theo."
), indent=True)

add_table(doc,
    "Bảng 3.4: Bảng ocr_sessions — Phiên nhận dạng ảnh Hán Nôm",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",          "INTEGER",   "PK, AUTOINCREMENT",  "Khóa chính tự tăng"),
        ("username",    "TEXT",      "NOT NULL",           "Email người dùng (tham chiếu Firebase)"),
        ("image_key",   "TEXT",      "NOT NULL",           "SHA256 hash của ảnh tài liệu"),
        ("image_name",  "TEXT",      "DEFAULT ''",         "Tên file ảnh gốc"),
        ("doc_name",    "TEXT",      "DEFAULT ''",         "Tên tài liệu (tự nhận dạng hoặc đặt tay)"),
        ("num_boxes",   "INTEGER",   "DEFAULT 0",          "Số vùng chữ trong ảnh"),
        ("is_favorite", "INTEGER",   "DEFAULT 0",          "Đánh dấu yêu thích (0/1)"),
        ("created_at",  "TIMESTAMP", "DEFAULT NOW",        "Thời gian tạo phiên"),
        ("updated_at",  "TIMESTAMP", "DEFAULT NOW",        "Thời gian cập nhật gần nhất"),
    ]
)

add_table(doc,
    "Bảng 3.5: Bảng ocr_boxes — Vùng chữ trong phiên OCR",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",            "INTEGER",   "PK, AUTOINCREMENT",       "Khóa chính tự tăng"),
        ("session_id",    "INTEGER",   "NOT NULL, FK → ocr_sessions.id", "Phiên OCR chứa vùng chữ này"),
        ("box_index",     "INTEGER",   "NOT NULL",                "Thứ tự vùng chữ trong trang (0 đến N-1)"),
        ("nom_ocr",       "TEXT",      "DEFAULT ''",              "Kết quả nhận dạng CRNN gốc"),
        ("nom_corrected", "TEXT",      "",                        "Chữ sau khi người dùng hiệu chỉnh"),
        ("hanviet",       "TEXT",      "DEFAULT ''",              "Phiên âm Hán Việt tra từ điển"),
        ("accuracy",      "TEXT",      "DEFAULT ''",              "Độ chính xác CRNN (VD: 85%)"),
        ("saved",         "INTEGER",   "DEFAULT 0",               "Đã lưu hiệu chỉnh hay chưa (0/1)"),
        ("updated_at",    "TIMESTAMP", "DEFAULT NOW",             "Thời gian cập nhật"),
    ]
)

# 3.3.2.3 Lịch sử
add_heading(doc, "3.3.2.3. Nhóm bảng Lịch sử", level=3)
add_body(doc, (
    "Nhóm này gồm 2 bảng ghi lại hoạt động người dùng, hoạt động độc lập và không "
    "có quan hệ FK với các nhóm bảng khác."
), indent=True)

add_table(doc,
    "Bảng 3.6: Bảng translation_log — Nhật ký dịch thuật trang chủ",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",          "INTEGER",   "PK, AUTOINCREMENT",  "Khóa chính tự tăng"),
        ("username",    "TEXT",      "NOT NULL",           "Email người dùng (tham chiếu Firebase)"),
        ("input_text",  "TEXT",      "",                   "Văn bản đầu vào"),
        ("output_text", "TEXT",      "",                   "Kết quả dịch"),
        ("direction",   "TEXT",      "",                   "Chiều dịch: vi_to_hn | hn_to_vi"),
        ("starred",     "INTEGER",   "DEFAULT 0",          "Đánh dấu yêu thích (0/1)"),
        ("created_at",  "TIMESTAMP", "DEFAULT NOW",        "Thời gian tra cứu"),
    ]
)

add_table(doc,
    "Bảng 3.7: Bảng corrections — Lưu hiệu chỉnh phục vụ tái huấn luyện",
    ["Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    [
        ("id",                "INTEGER",   "PK, AUTOINCREMENT", "Khóa chính tự tăng"),
        ("image_path",        "TEXT",      "",                  "Đường dẫn ảnh gốc"),
        ("original_ocr_text", "TEXT",      "",                  "Kết quả CRNN ban đầu"),
        ("corrected_text",    "TEXT",      "",                  "Văn bản sau khi người dùng hiệu chỉnh"),
        ("created_at",        "TIMESTAMP", "DEFAULT NOW",       "Thời gian tạo"),
    ]
)

# ── 3.3.3 Firebase ────────────────────────────────────
add_heading(doc, "3.3.3. Firebase (Cloud)", level=2)

add_heading(doc, "3.3.3.1. Firebase Authentication", level=3)
add_body(doc, (
    "Firebase Authentication là hệ thống quản lý danh tính người dùng của NomNaSite, "
    "tích hợp qua thư viện pyrebase4. Hệ thống hỗ trợ xác thực bằng email và mật khẩu. "
    "Chức năng quên mật khẩu được xử lý qua SMTP tự gửi mã OTP (mail_service.py), "
    "không dùng Firebase reset mặc định. Quản trị viên có thể khóa, mở khóa và xóa "
    "tài khoản trực tiếp trong ứng dụng qua Firebase Admin SDK (firebase_admin_service.py)."
), indent=True)

add_body(doc, "Các chức năng chính:")
add_body(doc, "•  Đăng nhập: sign_in_with_email_and_password() → trả về JWT token lưu vào session_state.")
add_body(doc, "•  Phân quyền: get_role(email) truy vấn Firebase Realtime DB → role admin | user.")
add_body(doc, "•  Quản lý tài khoản: khóa (disabled=True), mở khóa (disabled=False), xóa vĩnh viễn.")

add_heading(doc, "3.3.3.2. Firebase Realtime Database", level=3)
add_body(doc, (
    "Firebase Realtime Database được dùng riêng để lưu phân quyền người dùng, "
    "tách biệt hoàn toàn với dữ liệu OCR. Cấu trúc node đơn giản:"
), indent=True)
add_body(doc, '    roles/ { "<email_sanitized>": { "role": "admin" | "user" } }')
add_body(doc, "Email được sanitize (thay dấu '.' bằng ',') để làm key hợp lệ trong Firebase.")

# ── 3.3.4 Supabase ────────────────────────────────────
add_heading(doc, "3.3.4. Supabase (Cloud)", level=2)
add_body(doc, (
    "Supabase cung cấp Storage bucket trên cloud cho NomNaSite, dùng để lưu ảnh tài liệu "
    "OCR vĩnh viễn giữa các phiên làm việc. Khi người dùng nhấn 'Tiếp tục chỉnh sửa' "
    "trong trang Lịch sử, hệ thống tải ảnh từ Supabase về cache local để khôi phục "
    "phiên làm việc cũ. Nếu không cấu hình SUPABASE_URL và SUPABASE_KEY trong file .env, "
    "hệ thống tự động fallback sang SQLite local mà không báo lỗi."
), indent=True)

add_body(doc, "Cấu hình Storage:")
add_body(doc, "•  Bucket name: ocr-images")
add_body(doc, "•  Cấu trúc path: {username}/{image_key}.jpg")
add_body(doc, "•  Chức năng: upload_image() — lưu ảnh sau OCR; download_image() — tải về khi tiếp tục chỉnh sửa.")
add_body(doc, "•  Ngoài ra, khi Supabase được cấu hình, dữ liệu ocr_sessions và ocr_boxes được đồng bộ lên Supabase Database thay vì chỉ lưu SQLite local.")

# Lưu file
out = "e:/HOCTAP/ĐỒ ÁN TỐT NGHIỆP 2025/NomNaSite-main/Section_3_3_CSDL.docx"
doc.save(out)
print(f"Saved: {out}")
