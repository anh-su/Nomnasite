import base64
import cv2
import json
import os
import re
import time
import unicodedata
from pathlib import Path
import imagehash
import streamlit as st
import streamlit.components.v1 as components

from PIL import Image
from zipfile import ZipFile
from groq import Groq
from streamlit_drawable_canvas import st_canvas
from streamlit_javascript import st_javascript
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

from handler.asset import hash_bytes, download_assets, load_models, retrieve_image
from handler.bbox import generate_initial_drawing, transform_fabric_box, order_boxes4nom, get_patch
from handler.translator import db_hanviet, db_meaning, hvdic_render
from services import ocr_session as _ocr_svc
_ocr_svc.create_tables()
from toolbar import render_toolbar
from style import custom_css


_CSS_FILE  = Path(__file__).parent.parent / "css" / "nomnasite.css"
_IMGS_DIR  = Path(__file__).parent.parent / "imgs"
_BG_LIGHT  = base64.b64encode((_IMGS_DIR / "background.webp").read_bytes()).decode()

_ARCHIVE_DIR = Path(__file__).parent.parent / "data" / "archive" / "Raw"

# URL pattern → (tên, niên đại, url_page_regex)
_URL_PATTERNS = [
    (r'luc_van_tien/nlvnpf-\d+-(\d+)\.jpg',
     'Lục Vân Tiên', 'Giữa thế kỷ XIX (~1852)'),
    (r'kieu/1866/page(\d+)[ab]?\.jpg',
     'Truyện Kiều (bản 1866)', 'Đầu thế kỷ XIX (~1820), bản in 1866'),
    (r'kieu/1871/page(\d+)\.jpg',
     'Truyện Kiều (bản 1871)', 'Đầu thế kỷ XIX (~1820), bản in 1871'),
    (r'kieu/1872/page(\d+)[ab]?\.jpg',
     'Truyện Kiều (bản 1872)', 'Đầu thế kỷ XIX (~1820), bản in 1872'),
    (r'image_dvsk/quyen_thu/.*_(\d+)[ab]\.jpg',
     'ĐVSKTT – Quyển Thủ', 'Thế kỷ XV (1479), khắc in 1697'),
    (r'image_dvsk/ngoai_ky/.*_(\d+)[ab]\.jpg',
     'ĐVSKTT – Ngoại Kỷ Toàn Thư', 'Thế kỷ XV (1479), khắc in 1697'),
    (r'image_dvsk/ban_ky_toan_thu/.*_(\d+)[ab]\.jpg',
     'ĐVSKTT – Bản Kỷ Toàn Thư', 'Thế kỷ XV (1479), khắc in 1697'),
    (r'image_dvsk/ban_ky_thuc_luc/.*_(\d+)[ab]\.jpg',
     'ĐVSKTT – Bản Kỷ Thực Lục', 'Thế kỷ XV (1479), khắc in 1697'),
    (r'image_dvsk/ban_ky_tuc_bien/.*_(\d+)[ab]\.jpg',
     'ĐVSKTT – Bản Kỷ Tục Biên', 'Thế kỷ XVIII (1665–1675)'),
]

# Tên tài liệu → URL trang thông tin gốc trên nomfoundation.org
_DOC_INFO_URLS = {
    'Truyện Kiều (bản 1866)':       'https://nomfoundation.org/nom-project/tale-of-kieu/tale-of-kieu-version-1866?uiLang=en',
    'Truyện Kiều (bản 1871)':       'https://nomfoundation.org/nom-project/tale-of-kieu/tale-of-kieu-version-1871?uiLang=en',
    'Truyện Kiều (bản 1872)':       'https://nomfoundation.org/nom-project/tale-of-kieu/tale-of-kieu-version-1872?uiLang=en',
    'Lục Vân Tiên':                 'https://nomfoundation.org/nom-project/Luc-Van-Tien?uiLang=en',
    'Chinh phụ ngâm':               'https://nomfoundation.org/nom-project/Chinh-Phu-Ngam-Khuc?uiLang=en',
    'ĐVSKTT – Quyển Thủ':           'https://nomfoundation.org/nom-project/History-of-Greater-Vietnam?uiLang=en',
    'ĐVSKTT – Ngoại Kỷ Toàn Thư':  'https://nomfoundation.org/nom-project/History-of-Greater-Vietnam?uiLang=en',
    'ĐVSKTT – Bản Kỷ Toàn Thư':    'https://nomfoundation.org/nom-project/History-of-Greater-Vietnam?uiLang=en',
    'ĐVSKTT – Bản Kỷ Thực Lục':    'https://nomfoundation.org/nom-project/History-of-Greater-Vietnam?uiLang=en',
    'ĐVSKTT – Bản Kỷ Tục Biên':    'https://nomfoundation.org/nom-project/History-of-Greater-Vietnam?uiLang=en',
}

# Tên file (folder) → (tên, niên đại) để match khi upload
_FILENAME_TEXTS = {
    'luc van tien':      ('Lục Vân Tiên',             'Giữa thế kỷ XIX (~1852)'),
    'tale of kieu':      ('Truyện Kiều',               'Đầu thế kỷ XIX (~1820)'),
    'truyen kieu':       ('Truyện Kiều',               'Đầu thế kỷ XIX (~1820)'),
    'dvsktt':            ('Đại Việt Sử Ký Toàn Thư',  'Thế kỷ XV (1479), khắc in 1697'),
    'chinh phu ngam':    ('Chinh phụ ngâm',            'Thế kỷ XVIII (~1741)'),
    'cung oan ngam':     ('Cung oán ngâm khúc',        'Cuối thế kỷ XVIII (~1790)'),
}

# Build URL lookups từ url.txt
@st.cache_resource(show_spinner=False)
def _load_url_index():
    by_full: dict[str, str] = {}    # full_url  → folder_name
    by_base: dict[str, str] = {}    # basename  → full_url (first match)
    if not _ARCHIVE_DIR.exists():
        return by_full, by_base
    for url_file in _ARCHIVE_DIR.glob('*/url.txt'):
        folder = url_file.parent.name
        for line in url_file.read_text(encoding='utf-8').splitlines():
            line = line.strip()
            if not line:
                continue
            by_full[line] = folder
            base = line.rsplit('/', 1)[-1]
            if base not in by_base:
                by_base[base] = line
    return by_full, by_base

def _norm_title(s: str) -> str:
    s = unicodedata.normalize('NFD', s.lower())
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = re.sub(r'[^a-z0-9 ]', ' ', s)
    return re.sub(r'\s+', ' ', s).strip()

def _meta_from_url(u: str):
    """Trả về (ten, page, nien_dai) nếu URL khớp một pattern đã biết."""
    for pat, ten, nien_dai in _URL_PATTERNS:
        m = re.search(pat, u)
        if m:
            return ten, str(int(m.group(1))), nien_dai
    return None, '?', 'Không xác định'

def _lookup_doc_info(filename: str, url_input: str, n_boxes: int):
    dong = f'1 – {n_boxes}'
    by_full, by_base = _load_url_index()

    # 1. URL nhập tay — khớp chính xác hoặc theo pattern
    if url_input:
        u = url_input.strip()
        ten, page, nd = _meta_from_url(u)
        if ten:
            return ten, page, dong, u, nd
        # URL lạ nhưng có trong index
        folder = by_full.get(u)
        if folder:
            norm_f = _norm_title(folder)
            info   = next(((t, nd2) for k, (t, nd2) in _FILENAME_TEXTS.items() if k in norm_f), None)
            return (info[0] if info else folder), '?', dong, u, (info[1] if info else 'Không xác định')

    # 2. Upload file — khớp basename với url.txt để lấy URL gốc
    if filename:
        base = Path(filename).name
        src_url = by_base.get(base)
        if src_url:
            ten, page, nd = _meta_from_url(src_url)
            if ten:
                return ten, page, dong, src_url, nd

    # 3. Khớp tên file với bảng tên tác phẩm
    raw    = Path(filename).stem.replace('_', ' ').replace('-', ' ') if filename else ''
    page_m = re.search(r'\b[Pp]?(\d+)\s*$', raw)
    page   = page_m.group(1) if page_m else '1'
    title  = raw[:page_m.start()].strip() if page_m else raw.strip()
    norm   = _norm_title(title or raw)
    info   = next(((t, nd) for k, (t, nd) in _FILENAME_TEXTS.items() if k in norm), None)
    ten      = info[0] if info else (title or 'Không xác định')
    nien_dai = info[1] if info else 'Không xác định'
    src_url  = url_input or ''
    return ten, page, dong, src_url, nien_dai


_PAGES_DIR = Path(__file__).parent.parent / "data" / "archive" / "Pages"
_PHASH_THRESHOLD = 10

@st.cache_resource(show_spinner=False)
def _build_phash_index():
    """Build {phash: (basename, folder)} cho toàn bộ ảnh trong archive."""
    idx = {}
    if not _PAGES_DIR.exists():
        return idx
    for img_path in _PAGES_DIR.glob('*/imgs/*.jpg'):
        folder = img_path.parent.parent.name
        try:
            h = imagehash.phash(Image.open(img_path))
            idx[h] = (img_path.name, folder)
        except Exception:
            pass
    return idx

def _phash_lookup(pil_image: Image.Image):
    """Trả về (ten, page, nien_dai, src_url) nếu tìm thấy trong archive, else None."""
    idx = _build_phash_index()
    if not idx:
        return None
    try:
        h = imagehash.phash(pil_image)
    except Exception:
        return None
    best_dist, best_info = min(
        ((h - k, v) for k, v in idx.items()),
        key=lambda x: x[0]
    )
    if best_dist > _PHASH_THRESHOLD:
        return None
    basename, folder = best_info
    _, by_base = _load_url_index()
    src_url = by_base.get(basename, '')
    ten, page, nien_dai = _meta_from_url(src_url) if src_url else (folder, '?', 'Không xác định')
    return ten, page, nien_dai, src_url

def _ai_identify_doc(nom_text: str):
    """Dùng Groq để nhận diện tài liệu từ text OCR. Trả về (ten, nien_dai, nguon) hoặc None."""
    if not nom_text.strip():
        return None
    try:
        client = Groq(api_key=os.getenv('GROQ_API_KEY'))
        sample = nom_text[:300]
        resp = client.chat.completions.create(
            model='llama-3.3-70b-versatile',
            messages=[{
                'role': 'user',
                'content': (
                    f'Đây là đoạn văn bản chữ Nôm/Hán Nôm:\n{sample}\n\n'
                    'Hãy xác định tác phẩm này thuộc về tài liệu nào. '
                    'Trả lời ĐÚNG định dạng JSON sau và không có gì thêm:\n'
                    '{"ten":"...","nien_dai":"...","nguon":"..."}\n'
                    'Nếu không nhận ra thì để giá trị là chuỗi rỗng "".'
                )
            }],
            max_tokens=120,
            temperature=0.1,
        )
        raw = resp.choices[0].message.content.strip()
        m = re.search(r'\{.*\}', raw, re.DOTALL)
        if not m:
            return None
        data = json.loads(m.group())
        ten  = data.get('ten', '').strip()
        nd   = data.get('nien_dai', '').strip()
        src  = data.get('nguon', '').strip()
        if not ten:
            return None
        return ten, nd or 'Không xác định', src
    except Exception:
        return None


def _ai_find_doc_url(ten: str) -> str | None:
    """Dùng Groq tìm URL trang thông tin tài liệu, validate trước khi trả về."""
    if not ten or ten == 'Không xác định':
        return None
    try:
        import requests as _req
        client = Groq(api_key=os.getenv('GROQ_API_KEY'))
        resp = client.chat.completions.create(
            model='llama-3.3-70b-versatile',
            messages=[{
                'role': 'user',
                'content': (
                    f'Tài liệu Hán Nôm Việt Nam: "{ten}". '
                    'Cho tôi đúng một URL trang web học thuật hoặc thư viện số '
                    'giới thiệu tài liệu này (không phải link ảnh/JPG). '
                    'Chỉ trả về URL, không giải thích. Nếu không biết chắc, trả về chuỗi rỗng.'
                )
            }],
            max_tokens=80,
            temperature=0.0,
        )
        url = resp.choices[0].message.content.strip().strip('"\'')
        if not url or not url.startswith('http'):
            return None
        r = _req.head(url, timeout=5, allow_redirects=True)
        return url if r.status_code < 400 else None
    except Exception:
        return None


def show():
    st.markdown(custom_css, unsafe_allow_html=True)
    st.markdown(f"<style>{_CSS_FILE.read_text(encoding='utf-8')}</style>", unsafe_allow_html=True)
    st.markdown(f"""<style>
[data-testid="stAppViewContainer"] {{
    background-image: url("data:image/webp;base64,{_BG_LIGHT}");
    background-size: 25%;
    background-repeat: repeat;
    background-attachment: fixed;
}}
</style>""", unsafe_allow_html=True)
    det_model, rec_model = load_models()

    col1, col2 = st.columns(2)

    with col1:
        uploaded_file = st.file_uploader('Tải ảnh lên:', type=['jpg', 'jpeg', 'png'])
        url = st.text_input('Hoặc nhập URL ảnh:')
        image_path = retrieve_image(uploaded_file, url)

        # Chế độ tiếp tục phiên cũ: dùng ảnh đã lưu trong session_state hoặc Supabase
        if not image_path:
            _rkey = st.session_state.get('resume_image_key', '')
            if _rkey:
                image_path = st.session_state.get(f'img_path_{_rkey}', '')
                if not image_path and "username" in st.session_state:
                    with st.spinner('Đang tải lại ảnh...'):
                        _dl = _ocr_svc.download_image(st.session_state["username"], _rkey)
                    if _dl:
                        image_path = _dl
                        st.session_state[f'img_path_{_rkey}'] = _dl

        if not image_path or not os.path.exists(image_path):
            st.info('Vui lòng tải ảnh lên hoặc nhập URL ảnh.')
            st.session_state.pop('resume_image_key', None)
            return
        import numpy as _np
        _img_arr = cv2.imdecode(_np.fromfile(image_path, dtype=_np.uint8), cv2.IMREAD_COLOR)
        if _img_arr is None:
            st.warning('Không đọc được file ảnh. Vui lòng tải lại.')
            st.session_state.pop('resume_image_key', None)
            return
        raw_image = cv2.cvtColor(_img_arr, cv2.COLOR_BGR2RGB)
        _js_width = st_javascript('await fetch(window.location.href).then(response => window.innerWidth)')
        _col1_w = int(_js_width) if isinstance(_js_width, (int, float)) and _js_width > 100 else 700
        _img_w, _img_h = raw_image.shape[1], raw_image.shape[0]
        if _img_w <= _col1_w:
            canvas_width  = _img_w
            canvas_height = _img_h
        else:
            canvas_width  = _col1_w
            canvas_height = int(_img_h * _col1_w / _img_w)
        size_ratio = canvas_height / _img_h

        img_bytes = cv2.imencode('.jpg', raw_image)[1].tobytes()
        key = hash_bytes(img_bytes)

        # Lưu đường dẫn ảnh để dùng lại khi "Tiếp tục chỉnh sửa"
        st.session_state[f'img_path_{key}'] = image_path
        # Giữ resume_image_key — không pop, để rerun tự động vẫn tìm được ảnh

        # Upload ảnh lên Supabase (background thread, chỉ một lần mỗi ảnh)
        if "username" in st.session_state and not st.session_state.get(f'supa_up_{key}'):
            import threading as _threading
            _threading.Thread(
                target=_ocr_svc.upload_image,
                args=(st.session_state["username"], key, image_path),
                daemon=True,
            ).start()
            st.session_state[f'supa_up_{key}'] = True

        # Cache detection — chỉ chạy khi ảnh thay đổi, không chạy lại mỗi rerun
        if st.session_state.get('det_key') != key:
            with st.spinner('Đang phát hiện vùng chứa văn bản...'):
                boxes = det_model.predict_one_page(raw_image)
            st.session_state['det_key'] = key
            st.session_state['det_boxes'] = boxes
        else:
            boxes = st.session_state['det_boxes']

        # pHash lookup — sau detection để không chặn spinner
        if f'phash_{key}' not in st.session_state:
            st.session_state[f'phash_{key}'] = _phash_lookup(Image.fromarray(raw_image))

        # Khôi phục session_id từ DB — sau detection để không chặn spinner
        if "username" in st.session_state and f'sess_{key}' not in st.session_state:
            _existing_sess = _ocr_svc.find_session(st.session_state["username"], key)
            if _existing_sess:
                st.session_state[f'sess_{key}'] = _existing_sess

        mode, rec_clicked = render_toolbar(key)

        canvas_result = st_canvas(
            background_image = Image.open(image_path) if image_path else None,
            fill_color = 'rgba(76, 175, 80, 0.3)',
            width = max(canvas_width, 1),
            height = max(canvas_height, 1),
            stroke_width = 2,
            stroke_color = 'red',
            drawing_mode = 'rect' if mode == 'Vẽ' else 'transform',
            initial_drawing = generate_initial_drawing(boxes, size_ratio),
            update_streamlit = rec_clicked,
            key = f'canvas_{key}'
        )

    with col2:
        canvas_boxes = []
        if canvas_result.json_data and 'objects' in canvas_result.json_data:
            canvas_boxes = order_boxes4nom([
                transform_fabric_box(obj, size_ratio)
                for obj in canvas_result.json_data['objects']
            ])

        # Chỉ chạy OCR khi nhấn nút "Nhận dạng văn bản"
        if rec_clicked:
            with st.spinner('Đang nhận dạng văn bản trong từng khung...'):
                patches = []
                ocr_data = []
                _t0 = time.time()
                for box in canvas_boxes:
                    patch = get_patch(raw_image, box)
                    nom_text, confidence = rec_model.predict_one_patch(patch)
                    nom_text = nom_text.strip()
                    modern_text = db_hanviet(nom_text).strip()
                    points = sum(box.tolist(), [])
                    points_str = ','.join([str(round(p)) for p in points])
                    patches.append(patch)
                    ocr_data.append({
                        'nom': nom_text, 'modern': modern_text,
                        'accuracy': f'{confidence}%',
                        'points': points_str,
                        'height': str(patch.shape[0]), 'width': str(patch.shape[1])
                    })

                saved_json = {
                    'num_boxes': len(ocr_data),
                    'height': raw_image.shape[0],
                    'width': raw_image.shape[1],
                    'patches': ocr_data
                }
                with ZipFile('data/patches.zip', 'w') as zip_file:
                    with open('data/data.json', 'w', encoding='utf-8') as json_file, \
                         open('data/data.csv', 'w', encoding='utf-8', newline='') as csv_file:
                        csv_file.write('x1,y1,x2,y2,x3,y3,x4,y4,Chữ nôm,Nghĩa tiếng việt,chiều cao,chiều rộng\n')
                        for d, patch in zip(ocr_data, patches):
                            encoded = cv2.imencode('.jpg', cv2.cvtColor(patch, cv2.COLOR_BGR2RGB))[1]
                            zip_file.writestr(f'img/{d["nom"]}.jpg', encoded)
                            csv_file.write(f'{d["points"]},{d["nom"]},{d["modern"]},{d["height"]},{d["width"]}\n')
                        json.dump(saved_json, json_file, ensure_ascii=False, indent=4)
                    zip_file.write('data/data.json')
                    zip_file.write('data/data.csv')

            st.session_state['ocr_data'] = ocr_data
            st.session_state['ocr_boxes'] = canvas_boxes
            st.session_state['ocr_image_key'] = key
            st.session_state['ocr_time'] = round(time.time() - _t0, 2)

            # Tạo/cập nhật session lưu correction (chỉ khi đã login)
            if "username" in st.session_state:
                _img_name = uploaded_file.name if uploaded_file else (url or '')
                _new_sess_id = _ocr_svc.get_or_create_session(
                    st.session_state["username"], key, _img_name, num_boxes=len(ocr_data)
                )
                _ocr_svc.save_boxes(_new_sess_id, ocr_data)
                st.session_state[f'sess_{key}'] = _new_sess_id

            # AI fallback nếu pHash không nhận ra ảnh
            if st.session_state.get(f'phash_{key}') is None and f'ai_{key}' not in st.session_state:
                all_nom = ' '.join(d['nom'] for d in ocr_data if d['nom'])
                with st.spinner('Đang nhận diện tài liệu...'):
                    st.session_state[f'ai_{key}'] = _ai_identify_doc(all_nom)

        # Hiển thị kết quả từ session_state (giữ nguyên sau khi rerun)
        if st.session_state.get('ocr_image_key') == key:

            # Tải dữ liệu correction đã lưu từ DB
            _sess_id = st.session_state.get(f'sess_{key}')
            _saved_boxes = {}
            if _sess_id and "username" in st.session_state:
                _saved_boxes = _ocr_svc.get_boxes(_sess_id)

            for idx, (d, box) in enumerate(zip(
                st.session_state.get('ocr_data', []),
                st.session_state.get('ocr_boxes', [])
            )):
                patch = get_patch(raw_image, box)
                _binfo = _saved_boxes.get(idx, {})
                _is_saved = _binfo.get('saved', False)
                _exp_title = (
                    f':red[**Văn bản {idx + 1:02d}**:] {d["nom"]}' +
                    (' ✅' if _is_saved else '')
                )
                with st.expander(_exp_title):
                    col21, col22 = st.columns([1, 7])
                    with col21:
                        st.image(patch)
                    with col22:
                        conf_val = float(d.get('accuracy', '0%').rstrip('%'))
                        conf_color = '#27ae60' if conf_val >= 80 else ('#f39c12' if conf_val >= 50 else '#e74c3c')
                        acc = d.get('accuracy', 'N/A')
                        display = {
                            'Chữ Nôm': d['nom'],
                            'Phiên âm': d['modern'],
                            'Độ chính xác nhận dạng': (
                                f'🟢 {acc}' if conf_val >= 80
                                else f'🟡 {acc}' if conf_val >= 50
                                else f'🔴 {acc}'
                            ),
                            'Tọa độ': d['points'],
                            'Chiều cao': d['height'],
                            'Chiều rộng': d['width'],
                        }
                        st.table(display)
                    # Correction UI — chỉ hiện khi đã đăng nhập
                    if _sess_id and "username" in st.session_state:
                        st.markdown("---")
                        _cc1, _cc2 = st.columns([5, 1])
                        with _cc1:
                            _default_corr = _binfo.get('nom_corrected') or d['nom']
                            _corrected = st.text_input(
                                '✏️ Chỉnh sửa chữ Nôm:',
                                value=_default_corr,
                                key=f'corr_{_sess_id}_{idx}',
                            )
                        with _cc2:
                            st.markdown("<div style='padding-top:27px'></div>", unsafe_allow_html=True)
                            _btn_lbl = '✓ Lưu lại' if _is_saved else '💾 Lưu'
                            if st.button(_btn_lbl, key=f'savecorr_{_sess_id}_{idx}'):
                                _ocr_svc.save_correction(_sess_id, idx, _corrected)
                                # Lưu patch vào dataset để retrain sau này
                                try:
                                    from services.dataset_service import save_patch_from_correction
                                    save_patch_from_correction(
                                        patch, d['nom'], _corrected, conf_val / 100.0
                                    )
                                except Exception:
                                    pass
                                st.experimental_rerun()
                st.markdown(f'''
                    <b>Phiên âm:</b> {d["modern"]}<br/>
                    <b>Phiên âm</b> [hvdic](https://hvdic.thivien.net/transcript.php#trans): {hvdic_render(d["nom"])}<br/>
                    <b>Dịch nghĩa:</b> {db_meaning(d["nom"])}
                ''', unsafe_allow_html=True)

            # Chặn Streamlit auto-scroll col2 sau khi render xong, về lại đầu
            components.html("""
                <script>
                (function() {
                    var col2 = window.parent.document.querySelector(
                        '.block-container > div > [data-testid="stVerticalBlock"] > [data-testid="stHorizontalBlock"] > [data-testid="column"]:last-child'
                    );
                    if (!col2) return;
                    col2.addEventListener('scroll', function handler() {
                        col2.scrollTop = 0;
                        col2.removeEventListener('scroll', handler);
                    }, { once: true });
                    setTimeout(function() { col2.scrollTop = 0; }, 200);
                    setTimeout(function() { col2.scrollTop = 0; }, 700);
                })();
                </script>
            """, height=1)

    # ===== BANG SONG NGU (full-width, ben duoi hai cot) =====
    if st.session_state.get('ocr_image_key') == key:
        ocr_data = st.session_state.get('ocr_data', [])
        if ocr_data:
            st.markdown('<div style="font-size:18px;font-weight:700;color:#262660;margin:24px 0 12px 0;">1. Nội dung văn bản</div>', unsafe_allow_html=True)

            # -- Tên file export từ thông tin tài liệu --
            import json as _json, base64 as _b64
            from io import BytesIO as _BytesIO
            _pinfo = st.session_state.get(f'phash_{key}')
            _ainfo = st.session_state.get(f'ai_{key}')
            if _pinfo:
                _export_name = re.sub(r'[\\/*?:"<>|]', '_', f"{_pinfo[0]}_trang_{_pinfo[1]}")
            elif _ainfo:
                _export_name = re.sub(r'[\\/*?:"<>|]', '_', str(_ainfo[0]))
            else:
                _export_name = "han_nom_song_ngu"

            # -- Sinh nội dung 3 định dạng --
            def _build_bilingual():
                lines = []
                for _i, _d in enumerate(ocr_data, 1):
                    lines.append(_d["nom"])
                    lines.append(f"{_i}   {_d['modern']}")
                return "\n".join(lines)

            def _build_docx():
                from docx import Document as _Doc
                from docx.shared import Pt as _Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH as _WDA
                _doc = _Doc()
                for _i, _d in enumerate(ocr_data, 1):
                    _p1 = _doc.add_paragraph()
                    _p1.alignment = _WDA.CENTER
                    _p1.paragraph_format.space_before = _Pt(4)
                    _p1.paragraph_format.space_after  = _Pt(0)
                    _r1 = _p1.add_run(_d["nom"])
                    _r1.font.name = 'NomNaTong'
                    _r1.font.size = _Pt(18)
                    _p2 = _doc.add_paragraph()
                    _p2.alignment = _WDA.CENTER
                    _p2.paragraph_format.space_before = _Pt(0)
                    _p2.paragraph_format.space_after  = _Pt(6)
                    _p2.add_run(f"{_i}   {_d['modern']}").font.size = _Pt(13)
                _buf = _BytesIO()
                _doc.save(_buf)
                return _buf.getvalue()

            def _build_pdf():
                from fpdf import FPDF as _FPDF
                _fp = str(Path(__file__).parent.parent / "static" / "NomNaTong.otf")
                _pdf = _FPDF()
                _pdf.add_font("NomNaTong", "", _fp)
                _pdf.add_page()
                _pdf.set_auto_page_break(auto=True, margin=15)
                for _i, _d in enumerate(ocr_data, 1):
                    _pdf.set_font("NomNaTong", size=18)
                    _pdf.set_text_color(26, 26, 64)
                    _pdf.cell(0, 12, _d["nom"], new_x="LMARGIN", new_y="NEXT", align="C")
                    _pdf.set_font("NomNaTong", size=13)
                    _pdf.set_text_color(60, 60, 60)
                    _pdf.cell(0, 10, f"{_i}   {_d['modern']}", new_x="LMARGIN", new_y="NEXT", align="C")
                    _pdf.ln(2)
                return bytes(_pdf.output())

            _txt_b64  = _b64.b64encode(_build_bilingual().encode("utf-8")).decode()
            _docx_b64 = _b64.b64encode(_build_docx()).decode()
            _pdf_b64  = _b64.b64encode(_build_pdf()).decode()
            _fname_js = _json.dumps(_export_name)

            nom_lines  = "<br>".join(d["nom"] for d in ocr_data)
            viet_lines = "<br>".join(d["modern"] for d in ocr_data)
            nom_js  = _json.dumps("\n".join(d["nom"] for d in ocr_data))
            viet_js = _json.dumps("\n".join(d["modern"] for d in ocr_data))

            table_html = f"""
<style>
@font-face{{font-family:'NomNaTong';src:url('/app/static/NomNaTong.otf') format('opentype');}}
@font-face{{font-family:'HanaMin';src:url('/app/static/HanaMinA.otf') format('opentype');}}
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:transparent;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif}}
.nom-table{{width:100%;border-collapse:collapse;border-radius:12px;overflow:hidden;
           box-shadow:0 2px 16px rgba(38,38,96,0.10);}}
.nom-table th{{background:#262660;color:white;font-size:14px;font-weight:700;
              padding:12px 20px;text-align:left;letter-spacing:1.5px;}}
.nom-table th:first-child{{border-right:2px solid rgba(255,255,255,0.15);width:50%}}
.cpbtn{{float:right;background:rgba(255,255,255,0.15);border:1px solid rgba(255,255,255,0.3);
        color:white;border-radius:6px;padding:2px 10px;font-size:12px;cursor:pointer;
        transition:.15s;font-family:sans-serif;}}
.cpbtn:hover{{background:rgba(255,255,255,0.3);}}
.nom-td{{padding:16px 24px;font-size:20px;color:#1a1a40;border-right:2px solid #eee;
         font-family:'NomNaTong','Noto Sans CJK SC',serif;vertical-align:top;
         line-height:2;width:50%;text-align:center}}
.viet-td{{padding:16px 24px;font-size:15px;color:#333;vertical-align:top;
          line-height:2;width:50%;text-align:center}}
.dl-bar{{display:flex;justify-content:flex-end;gap:8px;padding:8px 2px 0 0;}}
.dlbtn{{background:#f4f6fb;border:1px solid #d0d7e8;color:#262660;border-radius:7px;
        padding:5px 14px;font-size:12px;font-weight:600;cursor:pointer;
        font-family:sans-serif;transition:.15s;}}
.dlbtn:hover{{background:#e8ecf8;border-color:#262660;}}
</style>
<script>
var _nomText  = {nom_js};
var _vietText = {viet_js};
var _fname    = {_fname_js};
var _files = {{
  txt:  {{b64:'{_txt_b64}',  mime:'text/plain', ext:'.txt'}},
  docx: {{b64:'{_docx_b64}', mime:'application/vnd.openxmlformats-officedocument.wordprocessingml.document', ext:'.docx'}},
  pdf:  {{b64:'{_pdf_b64}',  mime:'application/pdf', ext:'.pdf'}}
}};
function doCopy(btn,text){{
  navigator.clipboard.writeText(text).then(function(){{
    btn.textContent='✓ Đã sao chép';
    setTimeout(function(){{btn.textContent='📋 Sao chép';}},1500);
  }}).catch(function(){{
    var el=document.createElement('textarea');
    el.value=text;document.body.appendChild(el);el.select();
    document.execCommand('copy');document.body.removeChild(el);
    btn.textContent='✓ Đã sao chép';
    setTimeout(function(){{btn.textContent='📋 Sao chép';}},1500);
  }});
}}
function dlFile(type){{
  var f=_files[type];
  var bin=atob(f.b64);
  var arr=new Uint8Array(bin.length);
  for(var i=0;i<bin.length;i++)arr[i]=bin.charCodeAt(i);
  var blob=new Blob([arr],{{type:f.mime}});
  var a=document.createElement('a');
  a.href=URL.createObjectURL(blob);
  a.download=_fname+f.ext;
  a.click();
  URL.revokeObjectURL(a.href);
}}
</script>
<table class="nom-table">
  <thead>
    <tr>
      <th>漢 喃 &nbsp;·&nbsp; Hán Nôm &nbsp;<button class="cpbtn" onclick="doCopy(this,_nomText)">📋 Sao chép</button></th>
      <th>🇻🇳 &nbsp;Quốc Ngữ &nbsp;<button class="cpbtn" onclick="doCopy(this,_vietText)">📋 Sao chép</button></th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td class="nom-td">{nom_lines}</td>
      <td class="viet-td">{viet_lines}</td>
    </tr>
  </tbody>
</table>
<div class="dl-bar">
  <button class="dlbtn" onclick="dlFile('txt')">⬇ .txt</button>
  <button class="dlbtn" onclick="dlFile('docx')">⬇ .docx</button>
  <button class="dlbtn" onclick="dlFile('pdf')">⬇ .pdf</button>
</div>
"""
            height = max(300, len(ocr_data) * 44 + 100)
            components.html(table_html, height=height, scrolling=True)

            st.markdown('<div style="font-size:18px;font-weight:700;color:#262660;margin:6px 0 10px 0;">2. Thông tin nhận dạng</div>', unsafe_allow_html=True)

            so_vung     = len(ocr_data)
            so_chu      = sum(len(d["nom"]) for d in ocr_data)
            confidences = [float(d.get("accuracy","0%").rstrip("%")) for d in ocr_data]
            ti_le_cx    = round(sum(confidences) / len(confidences), 1) if confidences else 0
            ti_le_loi   = round(100 - ti_le_cx, 1)
            t_xl        = st.session_state.get("ocr_time", 0)

            metrics = [
                ("Số vùng ảnh",          f"{so_vung} vùng"),
                ("Số chữ nhận dạng",     f"{so_chu} chữ"),
                ("Tỉ lệ chính xác (TB)", f"{ti_le_cx}%"),
                ("Tỉ lệ lỗi (TB)",       f"{ti_le_loi}%"),
                ("Thời gian xử lý",      f"{t_xl}s"),
            ]
            cards_html = "".join(
                f'<div style="flex:1;background:white;border-radius:10px;padding:16px 10px;'
                f'text-align:center;box-shadow:0 1px 8px rgba(38,38,96,0.09);'
                f'border-top:3px solid #f69322;min-width:0;">'
                f'<div style="font-size:11px;color:#888;margin-bottom:8px;white-space:nowrap;">{lbl}</div>'
                f'<div style="font-size:22px;font-weight:700;color:#262660;">{val}</div>'
                f'</div>'
                for lbl, val in metrics
            )
            components.html(
                f'<div style="display:flex;gap:12px;font-family:sans-serif;">{cards_html}</div>',
                height=90
            )

            st.markdown('<div style="font-size:18px;font-weight:700;color:#262660;margin:18px 0 10px 0;">3. Thông tin tài liệu nguồn</div>', unsafe_allow_html=True)
            _phash_info = st.session_state.get(f'phash_{key}')
            _ai_info    = st.session_state.get(f'ai_{key}')

            if _phash_info:
                _ten, _trang, _nien_dai, _src_url = _phash_info
                _dong = f'1 – {len(ocr_data)}'
            elif _ai_info:
                _ten, _nien_dai, _src_url = _ai_info
                _trang = '?'
                _dong = f'1 – {len(ocr_data)}'
            else:
                _fname = uploaded_file.name if uploaded_file else ''
                _ten, _trang, _dong, _src_url, _nien_dai = _lookup_doc_info(_fname, url or '', len(ocr_data))

            # Cập nhật tên tài liệu vào session (một lần duy nhất)
            _sess_id_ref = st.session_state.get(f'sess_{key}')
            if (_sess_id_ref and _ten and _ten != 'Không xác định'
                    and not st.session_state.get(f'doc_upd_{_sess_id_ref}')):
                _ocr_svc.update_session_doc(_sess_id_ref, _ten)
                st.session_state[f'doc_upd_{_sess_id_ref}'] = True

            _web_url = _DOC_INFO_URLS.get(_ten)
            if not _web_url and _ten and _ten != 'Không xác định':
                _cache_key = f'doc_url_{_ten}'
                if _cache_key not in st.session_state:
                    with st.spinner('Đang tìm nguồn tài liệu...'):
                        st.session_state[_cache_key] = _ai_find_doc_url(_ten)
                _web_url = st.session_state[_cache_key]
            _link_url = _web_url or (_src_url if _src_url and _src_url.startswith('http') else None)
            if _link_url:
                _nguon_html = (
                    f'<a href="{_link_url}" target="_blank" '
                    f'style="color:#3355cc;text-decoration:none;word-break:break-all;">'
                    f'{_link_url}</a>'
                    f'&nbsp;<span style="color:#888;font-size:12px;">(trang {_trang})</span>'
                )
            elif _src_url:
                _nguon_html = f'<span style="color:#555;">{_src_url}</span>'
            else:
                _nguon_html = '<span style="color:#aaa;">Không xác định</span>'
            doc_fields_html = [
                ('Tên tài liệu', _ten),
                ('Trang',        _trang),
                ('Dòng',         _dong),
                ('Nguồn',        _nguon_html),
                ('Niên đại',     _nien_dai),
            ]
            doc_rows = "".join(
                f'<tr style="background:{"#f7f8ff" if i % 2 == 0 else "white"};">'
                f'<td style="padding:11px 20px;font-size:13px;font-weight:600;color:#5566aa;'
                f'width:170px;white-space:nowrap;border-left:3px solid #7b8fd4;">{lbl}</td>'
                f'<td style="padding:11px 20px;font-size:14px;color:#1a1a40;">{val}</td>'
                f'</tr>'
                for i, (lbl, val) in enumerate(doc_fields_html)
            )
            components.html(
                f'<div style="font-family:-apple-system,BlinkMacSystemFont,\'Segoe UI\',sans-serif;">'
                f'<table style="width:100%;border-collapse:collapse;border-radius:10px;overflow:hidden;'
                f'box-shadow:0 2px 12px rgba(38,38,96,0.10);">'
                f'{doc_rows}</table></div>',
                height=len(doc_fields_html) * 46 + 10
            )
