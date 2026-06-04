from datetime import datetime, timedelta, timezone
from pathlib import Path
import streamlit as st
import sqlite3
import pandas as pd
import os
from services import ocr_session as _ocr_svc
from style import bg_css

DB_PATH = "database/dictionary.db"
_CSS_FILE = Path(__file__).parent.parent / "css" / "history.css"


def get_connection():
    return sqlite3.connect(DB_PATH)




def _toggle_all(ids):
    val = st.session_state.get("sel_all_chk", False)
    for rid in ids:
        st.session_state[f"chk_{rid}"] = val


def conf_badge(conf):
    if conf is None:
        return ""
    pct = conf * 100
    if pct >= 80:
        cls, label = "conf-high",   f"✓ {pct:.0f}%"
    elif pct >= 50:
        cls, label = "conf-medium", f"~ {pct:.0f}%"
    else:
        cls, label = "conf-low",    f"✗ {pct:.0f}%"
    return f'<span class="conf-badge {cls}">{label}</span>'


def _safe_fname(s: str) -> str:
    import re
    return re.sub(r'[\\/*?:"<>|]', '_', s).strip() or 'export'


def _export_txt(boxes: list, title: str) -> bytes:
    """Xen kẽ: dòng chữ Nôm → dòng quốc ngữ (giống bảng song ngữ nomnasite)."""
    lines = [title, ''] if title else []
    for i, b in enumerate(boxes, 1):
        nom = b.get('nom_corrected') or b.get('nom_ocr') or ''
        hv  = b.get('hanviet') or ''
        lines.append(nom)
        lines.append(f"{i}   {hv}")
    return '\n'.join(lines).encode('utf-8')


def _export_docx(boxes: list, title: str) -> bytes:
    """Xen kẽ: đoạn chữ Nôm lớn → đoạn quốc ngữ nhỏ."""
    from io import BytesIO
    from docx import Document as _Doc
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WDA
    doc = _Doc()
    if title:
        h = doc.add_heading(title, 0)
        h.alignment = _WDA.CENTER
    for i, b in enumerate(boxes, 1):
        nom = b.get('nom_corrected') or b.get('nom_ocr') or ''
        hv  = b.get('hanviet') or ''
        p1 = doc.add_paragraph()
        p1.alignment = _WDA.CENTER
        p1.paragraph_format.space_before = _Pt(4)
        p1.paragraph_format.space_after  = _Pt(0)
        r1 = p1.add_run(nom)
        r1.font.name = 'NomNaTong'
        r1.font.size = _Pt(18)
        p2 = doc.add_paragraph()
        p2.alignment = _WDA.CENTER
        p2.paragraph_format.space_before = _Pt(0)
        p2.paragraph_format.space_after  = _Pt(6)
        p2.add_run(f"{i}   {hv}").font.size = _Pt(13)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_pdf(boxes: list, title: str) -> bytes:
    """Xen kẽ: dòng chữ Nôm lớn → dòng quốc ngữ nhỏ."""
    from fpdf import FPDF as _FPDF
    _fp = str(Path(__file__).parent.parent / 'static' / 'NomNaTong.otf')
    pdf = _FPDF()
    pdf.add_font('NomNaTong', '', _fp)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    if title:
        pdf.set_font('NomNaTong', size=16)
        pdf.set_text_color(26, 26, 64)
        pdf.cell(0, 12, title, new_x='LMARGIN', new_y='NEXT', align='C')
        pdf.ln(4)
    for i, b in enumerate(boxes, 1):
        nom = b.get('nom_corrected') or b.get('nom_ocr') or ''
        hv  = b.get('hanviet') or ''
        pdf.set_font('NomNaTong', size=18)
        pdf.set_text_color(26, 26, 64)
        pdf.cell(0, 12, nom, new_x='LMARGIN', new_y='NEXT', align='C')
        pdf.set_font('NomNaTong', size=13)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(0, 10, f"{i}   {hv}", new_x='LMARGIN', new_y='NEXT', align='C')
        pdf.ln(2)
    return bytes(pdf.output())


def _session_display_name(image_name: str, doc_name: str) -> tuple[str, str]:
    """Trả về (icon, label) thân thiện cho session title."""
    if doc_name:
        return '📖', doc_name
    if image_name:
        stem = image_name.rsplit('.', 1)[0] if '.' in image_name else image_name
        is_hash = len(stem) >= 32 and all(c in '0123456789abcdef' for c in stem.lower())
        if not is_hash:
            return '📄', image_name
    return '🖼️', 'Ảnh không tên'


def _parse_dt(s: str):
    if not s:
        return None
    try:
        s = s.replace(' ', 'T')
        if '+' not in s and 'Z' not in s:
            s += '+00:00'
        return datetime.fromisoformat(s)
    except Exception:
        return None


def _fmt_dt(s: str) -> str:
    """ISO timestamp → giờ Việt Nam (UTC+7), format DD/MM/YYYY HH:MM."""
    dt = _parse_dt(s)
    if not dt:
        return (s or '')[:16].replace('T', ' ')
    dt_vn = dt + timedelta(hours=7)
    return dt_vn.strftime('%d/%m/%Y %H:%M')


def _apply_filter(sessions: list, f: str) -> list:
    if f == 'Tất cả':
        return sessions
    if f == 'Yêu thích':
        return [s for s in sessions if s.get('is_favorite')]
    now = datetime.now(timezone.utc)
    if f == 'Hôm nay':
        cutoff = now.replace(hour=0, minute=0, second=0, microsecond=0)
    else:  # Tuần này
        cutoff = (now - timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
    result = []
    for s in sessions:
        dt = _parse_dt(s.get('updated_at') or s.get('created_at', ''))
        if dt and dt >= cutoff:
            result.append(s)
    return result


def _render_ocr_sessions(username):
    # Xử lý pending resume TRƯỚC khi render expanders
    _pending = st.session_state.pop('_resume_pending', None)
    if _pending:
        st.session_state['page'] = 'nomnasite'
        st.session_state['resume_image_key'] = _pending['image_key']
        if _pending.get('user') and _pending.get('name'):
            st.experimental_set_query_params(
                user=_pending['user'],
                name=_pending['name'],
                page='nomnasite',
            )
        else:
            st.experimental_set_query_params(page='nomnasite')
        st.experimental_rerun()

    sessions = _ocr_svc.get_sessions(username, limit=50)

    if not sessions:
        st.markdown("""
        <div class="empty-state">
            <div class="empty-icon">📭</div>
            <p>Chưa có phiên OCR nào được lưu.</p>
            <p style="font-size:13px">Hãy dùng <b>Nhận diện hình ảnh</b>, nhận dạng xong rồi lưu correction.</p>
        </div>
        """, unsafe_allow_html=True)
        return

    # ── Bộ lọc ───────────────────────────────────────────────────────────────
    n_fav = sum(1 for s in sessions if s.get('is_favorite'))
    _filter_opts = ['Tất cả', 'Hôm nay', 'Tuần này', f'Yêu thích ({n_fav})']
    _filter_keys = ['Tất cả', 'Hôm nay', 'Tuần này', 'Yêu thích']
    _fi = st.radio(
        'Lọc:', _filter_opts, horizontal=True,
        key='hist_filter', label_visibility='collapsed'
    )
    active_filter = _filter_keys[_filter_opts.index(_fi)]
    sessions = _apply_filter(sessions, active_filter)

    # Thống kê tổng quan
    total_sess  = len(sessions)
    total_saved = sum(s['saved_boxes'] for s in sessions)
    total_boxes = sum(s['total_boxes'] for s in sessions)

    if total_sess == 0:
        st.markdown("""
        <div class="empty-state">
            <div class="empty-icon">🔍</div>
            <p>Không có phiên nào khớp với bộ lọc này.</p>
        </div>
        """, unsafe_allow_html=True)
        return

    st.markdown(f"""
    <div class="stat-grid">
        <div class="stat-card">
            <div class="stat-icon">🖼️</div>
            <div>
                <div class="stat-val">{total_sess}</div>
                <div class="stat-lbl">Phiên OCR</div>
            </div>
        </div>
        <div class="stat-card">
            <div class="stat-icon">✅</div>
            <div>
                <div class="stat-val">{total_saved}/{total_boxes}</div>
                <div class="stat-lbl">Văn bản đã lưu</div>
            </div>
        </div>
        <div class="stat-card">
            <div class="stat-icon">⭐</div>
            <div>
                <div class="stat-val">{n_fav}</div>
                <div class="stat-lbl">Yêu thích</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    for sess in sessions:
        _sid    = sess['id']
        _doc    = sess['doc_name'] or ''
        _date   = _fmt_dt(sess['updated_at'] or sess['created_at'] or '')
        _total  = sess['total_boxes']
        _saved  = sess['saved_boxes']
        _pct    = int(_saved / _total * 100) if _total > 0 else 0
        _filled = round(_pct / 10)
        _bar    = '█' * _filled + '░' * (10 - _filled)
        _prog_color = '#27ae60' if _pct == 100 else ('#f39c12' if _pct > 0 else '#aaa')
        _ico, _label = _session_display_name(sess['image_name'], _doc)
        _is_fav  = sess.get('is_favorite', False)
        _fav_tag = ' ⭐' if _is_fav else ''
        _prog_tag = f' &nbsp;·&nbsp; ✅ {_saved}/{_total}' if _total > 0 else ''
        _exp_label = f'{_ico} **{_label}**{_fav_tag} &nbsp;|&nbsp; 🕒 {_date}{_prog_tag}'

        with st.expander(_exp_label):
            # Progress bar
            st.markdown(f"""
<div style="margin-bottom:12px;font-family:sans-serif">
  <span style="font-size:12px;color:#888">Tiến độ chỉnh sửa: </span>
  <span style="font-weight:700;color:{_prog_color}">{_saved}/{_total}</span>
  <code style="font-family:monospace;color:{_prog_color};margin-left:6px">[{_bar}] {_pct}%</code>
</div>
""", unsafe_allow_html=True)

            # Danh sách boxes
            boxes = _ocr_svc.get_session_boxes(_sid)
            if boxes:
                for b in boxes:
                    _bidx    = b['box_index'] + 1
                    _nom_ocr = b['nom_ocr'] or '—'
                    _nom_cor = b['nom_corrected']
                    _hv      = b['hanviet'] or '—'
                    _is_sv   = b['saved']
                    _icon    = '✅' if _is_sv else '⬜'
                    _lbl_clr = '#27ae60' if _is_sv else '#999'
                    _lbl_bg  = '#e8f5e9' if _is_sv else '#f5f5f5'
                    _lbl_txt = 'Đã lưu' if _is_sv else 'Chưa lưu'
                    _cor_html = (
                        f' &nbsp;→&nbsp; <b style="color:#262660">{_nom_cor}</b>'
                        if _nom_cor and _nom_cor != _nom_ocr else ''
                    )
                    st.markdown(f"""
<div style="display:flex;align-items:center;gap:10px;padding:6px 2px;
            border-bottom:1px solid #f0f0f0;font-family:sans-serif;font-size:13px;">
  <span style="flex-shrink:0;font-size:15px">{_icon}</span>
  <span style="min-width:85px;color:#777;flex-shrink:0">Văn bản {_bidx:02d}</span>
  <span style="font-family:'NomNaTong',serif;font-size:17px;color:#1a1a40;flex:1">
    {_nom_ocr}{_cor_html}
  </span>
  <span style="color:#aaa;flex-shrink:0;font-size:12px">{_hv}</span>
  <span style="color:{_lbl_clr};font-size:11px;flex-shrink:0;background:{_lbl_bg};
               padding:2px 8px;border-radius:10px;">{_lbl_txt}</span>
</div>
""", unsafe_allow_html=True)

            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

            # Lấy các box đã lưu để export
            _saved_boxes = [b for b in (boxes or []) if b.get('saved')]
            _export_title = _label
            _fname = _safe_fname(_export_title)
            _has_exp = len(_saved_boxes) > 0

            _ac1, _ac_t, _ac_d, _ac_p, _ac2, _ac3 = st.columns([2.5, 1, 1, 1, 0.6, 0.6])
            with _ac1:
                if st.button('🔄 Tiếp tục chỉnh sửa', key=f'resume_{_sid}', use_container_width=True):
                    st.session_state['_resume_pending'] = {
                        'image_key': sess.get('image_key', ''),
                        'user': st.session_state.get('user', ''),
                        'name': st.session_state.get('username', ''),
                    }
                    st.experimental_rerun()
            with _ac_t:
                if _has_exp:
                    st.download_button(
                        '⬇ .txt',
                        data=_export_txt(_saved_boxes, _export_title),
                        file_name=f'{_fname}.txt',
                        mime='text/plain',
                        key=f'dl_txt_{_sid}',
                        use_container_width=True,
                    )
                else:
                    st.button('⬇ .txt', key=f'dl_txt_{_sid}', disabled=True, use_container_width=True)
            with _ac_d:
                if _has_exp:
                    st.download_button(
                        '⬇ .docx',
                        data=_export_docx(_saved_boxes, _export_title),
                        file_name=f'{_fname}.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        key=f'dl_doc_{_sid}',
                        use_container_width=True,
                    )
                else:
                    st.button('⬇ .docx', key=f'dl_doc_{_sid}', disabled=True, use_container_width=True)
            with _ac_p:
                if _has_exp:
                    st.download_button(
                        '⬇ .pdf',
                        data=_export_pdf(_saved_boxes, _export_title),
                        file_name=f'{_fname}.pdf',
                        mime='application/pdf',
                        key=f'dl_pdf_{_sid}',
                        use_container_width=True,
                    )
                else:
                    st.button('⬇ .pdf', key=f'dl_pdf_{_sid}', disabled=True, use_container_width=True)
            with _ac2:
                _star_lbl = '⭐' if _is_fav else '☆'
                if st.button(_star_lbl, key=f'fav_{_sid}', help='Đánh dấu yêu thích', use_container_width=True):
                    _ocr_svc.toggle_favorite(_sid)
                    st.experimental_rerun()
            with _ac3:
                if st.button('🗑️', key=f'delsess_{_sid}', help='Xóa phiên', use_container_width=True):
                    st.session_state[f'_confirm_sess_{_sid}'] = True
                    st.experimental_rerun()

            if st.session_state.get(f'_confirm_sess_{_sid}'):
                st.warning('Xóa phiên này? Tất cả correction sẽ mất, không hoàn tác được!')
                _cy, _cn, _ = st.columns([1.2, 1, 5])
                with _cy:
                    if st.button('✓ Xác nhận', key=f'yes_sess_{_sid}'):
                        _ocr_svc.delete_session(_sid)
                        st.session_state.pop(f'_confirm_sess_{_sid}', None)
                        st.experimental_rerun()
                with _cn:
                    if st.button('✗ Hủy', key=f'no_sess_{_sid}'):
                        st.session_state.pop(f'_confirm_sess_{_sid}', None)
                        st.experimental_rerun()




def show():
    if "username" not in st.session_state:
        st.warning("Vui lòng đăng nhập để xem lịch sử.")
        return

    _ocr_svc.create_tables()

    if _CSS_FILE.exists():
        st.markdown(f"<style>{_CSS_FILE.read_text(encoding='utf-8')}</style>",
                    unsafe_allow_html=True)
    st.markdown(bg_css(), unsafe_allow_html=True)

    username = st.session_state["username"]

    st.markdown(f"""
    <div class="hist-header">
        <h2>📜 Lịch sử nhận dạng OCR</h2>
        <p>Xem lại các phiên dịch ảnh chữ Hán Nôm của <b>{username}</b></p>
    </div>
    """, unsafe_allow_html=True)

    _render_ocr_sessions(username)
